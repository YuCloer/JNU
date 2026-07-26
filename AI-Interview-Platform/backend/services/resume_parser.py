"""简历解析服务：分段提取 + 极简prompt + 强后处理"""
import json
import re

from pydantic import ValidationError

from schemas import ResumeSchema
from services.llm_client import llm_json_strict as llm_json

# ===== 极简 prompt（3B模型只能跟住短指令）=====

EXTRACT_PROMPT = """从简历文本提取信息，返回JSON。

示例输入：
张三
手机：13800001111 邮箱：zhang@qq.com
教育背景
2021.09-2025.06 武汉大学 计算机科学 本科
技能：Python, SQL, Git
语言能力
英语 CET-6
法语 DELF B2
实习经历
2024.07-2024.09 字节跳动 数据分析师实习生
负责用户行为数据看板搭建
项目经历
AI简历分析平台 负责人
技术栈：FastAPI, Vue3, Ollama

示例输出：
{{"name":"张三","email":"zhang@qq.com","phone":"13800001111","education":[{{"school":"武汉大学","major":"计算机科学","degree":"本科","start_date":"2021.09","end_date":"2025.06"}}],"skills":["Python","SQL","Git"],"languages":["英语 CET-6","法语 DELF B2"],"experiences":[{{"company":"字节跳动","position":"数据分析师实习生","duration":"2024.07-2024.09","description":"负责用户行为数据看板搭建"}}],"projects":[{{"name":"AI简历分析平台","role":"负责人","description":"","tech_stack":"FastAPI, Vue3, Ollama"}}]}}

现在提取以下简历，只返回JSON：
{text}"""


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """根据文件类型提取纯文本"""
    if filename.endswith(".pdf"):
        return _extract_pdf(file_content)
    elif filename.endswith(".docx"):
        return _extract_docx(file_content)
    else:
        raise ValueError(f"不支持的文件格式: {filename}")


def _extract_pdf(content: bytes) -> str:
    """用 pdfplumber 提取PDF文本，自动检测多栏布局并按栏顺序提取"""
    import io
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = _extract_page_columns(page)
                text += page_text + "\n"
            if text.strip():
                return text
    except ImportError:
        pass
    # fallback
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(content))
    text = ""
    for page in reader.pages:
        text += (page.extract_text() or "") + "\n"
    return text


def _extract_page_columns(page) -> str:
    """检测页面是否为多栏布局，是则按左→右顺序提取，否则正常提取"""
    page_width = page.width
    words = page.extract_words()

    if not words:
        return page.extract_text() or ""

    # 检测分栏：在页面中间区域(30%-70%)找垂直空白带
    # 统计每个x位置区间内的word数量
    mid_left = page_width * 0.25
    mid_right = page_width * 0.75
    bin_size = page_width / 50  # 把页面分成50个bin
    bins = [0] * 51

    for w in words:
        # word的x0和x1取中点
        x_mid = (w["x0"] + w["x1"]) / 2
        bin_idx = min(int(x_mid / bin_size), 50)
        bins[bin_idx] += 1

    # 在25%-75%区间找最宽的空白带（连续bin为0或极少）
    best_gap_start = -1
    best_gap_width = 0
    gap_start = -1
    threshold = max(1, len(words) * 0.005)  # 极少word的bin视为空白

    for i in range(int(50 * 0.25), int(50 * 0.75)):
        if bins[i] <= threshold:
            if gap_start == -1:
                gap_start = i
        else:
            if gap_start != -1:
                gap_width = i - gap_start
                if gap_width > best_gap_width:
                    best_gap_width = gap_width
                    best_gap_start = gap_start
                gap_start = -1
    # 处理尾部
    if gap_start != -1:
        gap_width = int(50 * 0.75) - gap_start
        if gap_width > best_gap_width:
            best_gap_width = gap_width
            best_gap_start = gap_start

    # 空白带宽度至少占页面3%才算分栏
    if best_gap_width >= 2 and best_gap_start > 0:
        split_x = (best_gap_start + best_gap_width / 2) * bin_size
        # 分栏提取
        left = page.crop((0, 0, split_x, page.height))
        right = page.crop((split_x, 0, page_width, page.height))
        left_text = left.extract_text() or ""
        right_text = right.extract_text() or ""
        # 宽栏（主内容）在前，窄栏（侧边栏）在后
        # 判断哪边是主内容：文字更多的一边
        if len(right_text) >= len(left_text):
            return right_text + "\n" + left_text
        else:
            return left_text + "\n" + right_text
    else:
        # 单栏，正常提取
        return page.extract_text() or ""


def _extract_docx(content: bytes) -> str:
    import docx
    import io
    doc = docx.Document(io.BytesIO(content))
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
    # 也提取表格内容（很多简历用表格排版）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" ".join(cells))
    return "\n".join(paragraphs)


def _sanitize_llm_output(data: dict) -> dict:
    """修正LLM输出的格式问题，避免pydantic验证失败导致整个结果丢弃"""
    # languages: 模型可能输出 [{language, level, ...}] 而非 ["英语 CET-6"]
    if "languages" in data and isinstance(data["languages"], list):
        fixed = []
        for item in data["languages"]:
            if isinstance(item, str):
                fixed.append(item)
            elif isinstance(item, dict):
                lang = item.get("language", "")
                level = item.get("level", "") or item.get("certification", "")
                if lang and level:
                    fixed.append(f"{lang} {level}")
                elif lang:
                    fixed.append(lang)
                elif level:
                    fixed.append(level)
        data["languages"] = fixed

    # education: 确保是 list[dict]，每项有正确字段
    if "education" in data and isinstance(data["education"], list):
        for edu in data["education"]:
            if isinstance(edu, dict):
                # start_date: 只保留年份部分，去掉"至今"等后缀
                sd = edu.get("start_date", "")
                if sd:
                    ym = re.search(r"(20\d{2})[./\-年]?(\d{1,2})?", sd)
                    if ym:
                        edu["start_date"] = ym.group(1) + ("." + ym.group(2).zfill(2) if ym.group(2) else "")
                    elif not re.search(r"20\d{2}", sd):
                        edu["start_date"] = ""
                # end_date: 规范化
                ed = edu.get("end_date", "")
                if ed:
                    if re.search(r"至今|在读|present|今|now", ed, re.IGNORECASE):
                        edu["end_date"] = "至今"
                    else:
                        ym = re.search(r"(20\d{2})\s*届?", ed)
                        if ym:
                            # "2027 届" → "2027"，"2025.06" → "2025.06"
                            full = re.search(r"(20\d{2})[./\-年](\d{1,2})", ed)
                            if full:
                                edu["end_date"] = full.group(1) + "." + full.group(2).zfill(2)
                            else:
                                edu["end_date"] = ym.group(1)
                        else:
                            edu["end_date"] = ""

    # experiences: 模型可能输出 internships 而非 experiences
    if "internships" in data and "experiences" not in data:
        data["experiences"] = data.pop("internships")

    # projects: 确保 tech_stack 是字符串
    if "projects" in data and isinstance(data["projects"], list):
        for proj in data["projects"]:
            if isinstance(proj, dict) and isinstance(proj.get("tech_stack"), list):
                # 3B模型可能输出 [{"name":"Python"}] 而非 ["Python"]
                fixed_ts = []
                for ts_item in proj["tech_stack"]:
                    if isinstance(ts_item, str):
                        fixed_ts.append(ts_item)
                    elif isinstance(ts_item, dict):
                        name = ts_item.get("name", "") or ts_item.get("tech", "")
                        if name:
                            fixed_ts.append(name)
                proj["tech_stack"] = ", ".join(fixed_ts)
            # role 字段不应包含GitHub路径或斜杠路径
            if isinstance(proj, dict):
                role = proj.get("role", "")
                if role and re.search(r"github\.com|/[\w-]+/[\w-]+|^[\w-]+/[\w-]+/[\w-]+$", role):
                    proj["role"] = ""

    # skills: 模型可能输出 [{name, level}] 而非 ["Python", "SQL"]
    if "skills" in data and isinstance(data["skills"], list):
        fixed = []
        for item in data["skills"]:
            if isinstance(item, str):
                fixed.append(item)
            elif isinstance(item, dict):
                name = item.get("name", "") or item.get("skill", "")
                if name:
                    fixed.append(name)
        data["skills"] = fixed

    # experiences: 确保每项的字段是字符串
    if "experiences" in data and isinstance(data["experiences"], list):
        for exp in data["experiences"]:
            if isinstance(exp, dict):
                for key in ("company", "position", "duration", "description"):
                    val = exp.get(key)
                    if isinstance(val, list):
                        exp[key] = "；".join(str(v) for v in val)

    return data


def extract_resume(raw_text: str) -> dict:
    """解析简历：LLM提取 → 后处理修正 → 正则兜底"""
    # 安全网：预处理前先从原始文本提取联系方式（防止预处理丢失）
    raw_email = _regex_email(raw_text)
    raw_phone = _regex_phone(raw_text)

    # 预处理：清理多余空白
    cleaned = _preprocess_text(raw_text)

    # 第一层：LLM
    parsed = None
    try:
        result = llm_json.invoke(EXTRACT_PROMPT.format(text=cleaned[:3500]))
        data = json.loads(result.content)
        data = _sanitize_llm_output(data)
        validated = ResumeSchema(**data)
        parsed = validated.model_dump()
    except Exception:
        pass

    # 第二层：正则兜底
    if parsed is None:
        parsed = {
            "name": "",
            "email": _regex_email(cleaned),
            "phone": _regex_phone(cleaned),
            "education": [],
            "skills": [],
            "experiences": [],
            "projects": [],
        }

    # ===== 强后处理：用正则校验/补全每个字段 =====
    parsed["name"] = _ensure_name(parsed.get("name", ""), cleaned)
    parsed["email"] = parsed.get("email") or _regex_email(cleaned) or raw_email
    parsed["phone"] = parsed.get("phone") or _regex_phone(cleaned) or raw_phone

    # 教育：LLM没提取到就用正则补
    if not parsed.get("education"):
        parsed["education"] = _regex_education(cleaned)

    # 语言能力：验证LLM输出质量，不合格就用正则
    llm_langs = parsed.get("languages") or []
    # 检查LLM输出是否包含真正的语言名
    _valid_lang_pattern = re.compile(r"(英语|英文|English|法语|French|日语|Japanese|韩语|Korean|德语|German|中文|普通话|粤语)", re.IGNORECASE)
    valid_langs = [l for l in llm_langs if isinstance(l, str) and _valid_lang_pattern.search(l)]
    if valid_langs:
        parsed["languages"] = valid_langs
    else:
        parsed["languages"] = _regex_languages(cleaned)
    # 如果正则也没找到，尝试从原始文本提取
    if not parsed.get("languages"):
        parsed["languages"] = _regex_languages(raw_text)

    # 技能：LLM提取 + 正则补全，统一过滤
    llm_skills = parsed.get("skills") or []
    regex_skills = _regex_skills(cleaned)
    # 合并去重（LLM优先，正则补充）
    merged = list(llm_skills)
    for s in regex_skills:
        if s not in merged:
            merged.append(s)
    parsed["skills"] = _filter_skills(merged)

    # 课程推断技能：学了某门课 = 具备相关知识
    course_skills = _courses_to_skills(cleaned)
    for s in course_skills:
        if s not in parsed["skills"]:
            parsed["skills"].append(s)

    # 合并相似技能（用/隔开）
    parsed["skills"] = _merge_similar_skills(parsed["skills"])

    # 经历：LLM没提取到就用正则补
    if not parsed.get("experiences"):
        parsed["experiences"] = _regex_experiences(cleaned)

    # 项目：LLM没提取到就用正则补
    if not parsed.get("projects"):
        parsed["projects"] = _regex_projects(cleaned)

    return parsed


def _preprocess_text(text: str) -> str:
    """清理PDF提取的乱序空白和多栏混排，重组断行的核心课程"""
    # 去掉连续空格（保留单个）
    text = re.sub(r"[ \t]{2,}", " ", text)
    # 去掉空行堆积
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    # === 重组被多栏截断的核心课程 ===
    text = _reassemble_courses(text)

    # === 清理"开源项目"中的GitHub路径链接 ===
    # 侧边栏常有 "开源项目\nGitHub / User / Org / Repo-\nName" 这种格式
    # LLM会误认为是项目名，需要合并断行并标记为链接
    text = _clean_github_links(text)

    return text


def _clean_github_links(text: str) -> str:
    """合并断行的GitHub路径，标记为链接而非项目"""
    lines = text.split("\n")
    result = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        # 检测"开源项目"段落标题
        if re.match(r"^开源项[⽬目]$", stripped):
            result.append("开源项目（GitHub仓库）")
            i += 1
            # 收集后续的GitHub路径行，合并断行
            github_buffer = ""
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue
                # GitHub路径行特征：含 "GitHub" 或 "/" 分隔的路径
                if "GitHub" in line or (github_buffer and "/" in line):
                    github_buffer += (" " if github_buffer else "") + line
                    # 检查是否路径结束（下一行不含 "/" 且不以 "-" 结尾）
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if not next_line.startswith(("GitHub", "-")) and "/" not in next_line and not github_buffer.rstrip().endswith("-"):
                            # 路径结束，输出合并后的行
                            # 去掉 "GitHub / " 前缀，只保留仓库名
                            repo = re.sub(r"^GitHub\s*/\s*", "", github_buffer)
                            repo = repo.replace(" ", "").replace("-", "", 1) if repo.endswith("-") else repo.replace(" ", "")
                            # 修复断行：如 "Grade-\nWatcher" → "Grade-Watcher"
                            repo = re.sub(r"-\s*$", "-", repo)
                            result.append(f"  仓库: {repo}")
                            github_buffer = ""
                    i += 1
                elif github_buffer and (github_buffer.rstrip().endswith("-") or "/" in line):
                    # 断行续接（如 "Grade-" 后接 "Watcher"）
                    github_buffer += line
                    i += 1
                    # 续接完成
                    repo = re.sub(r"^GitHub\s*/\s*", "", github_buffer)
                    repo = repo.replace(" / ", "/").replace(" ", "")
                    result.append(f"  仓库: {repo}")
                    github_buffer = ""
                else:
                    break
            if github_buffer:
                repo = re.sub(r"^GitHub\s*/\s*", "", github_buffer)
                repo = repo.replace(" / ", "/").replace(" ", "")
                result.append(f"  仓库: {repo}")
            continue
        result.append(lines[i])
        i += 1

    # 第二遍：处理散落在其他位置的 github.com URL 格式
    # 将 "https://github.com/user/repo" 或 "github.com/user/repo" 标记为仓库链接
    # 也处理纯 "User/Repo" 或 "User/Org/Repo" 格式（侧边栏常见）
    final = []
    for line in result:
        stripped = line.strip()
        m = re.match(r"^(?:https?://)?github\.com/([\w.-]+/[\w.-]+(?:/[\w.-]+)?)/?\s*$", stripped)
        if m:
            final.append(f"  仓库: {m.group(1)}")
        elif re.fullmatch(r"[A-Za-z0-9_.-]+/[A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)?", stripped) and not stripped.startswith("/"):
            # 纯 User/Repo 或 User/Org/Repo 格式（排除文件路径如 src/main）
            parts = stripped.split("/")
            if all(p[0].isupper() or p[0].isdigit() for p in parts) or any(c.isupper() for c in stripped):
                final.append(f"  仓库: {stripped}")
            else:
                final.append(line)
        else:
            final.append(line)
    return "\n".join(final)


def _reassemble_courses(text: str) -> str:
    """检测并重组被多栏截断的核心课程列表"""
    lines = text.split("\n")
    result_lines = []
    course_buffer = []
    contact_lines = []  # 保存被课程段落截断的联系方式行
    in_courses = False

    for line in lines:
        stripped = line.strip()

        # 检测课程段落开始
        if re.search(r"核[心⼼]课程", stripped):
            in_courses = True
            course_buffer = [stripped]
            continue

        if in_courses:
            # 判断是否是纯噪音行（只有邮箱/电话，没有课程信息）
            is_pure_noise = bool(re.search(
                r"^(邮箱|电话|手机|email|phone)\s", stripped, re.IGNORECASE
            )) and not re.search(r"(程序设计|系统|原理|分析|设计|学|论|基础|·)", stripped)

            # 也检测纯邮箱/纯电话行（无前缀）
            is_bare_contact = bool(
                re.fullmatch(r"[\w.+-]+@[\w-]+\.[\w.-]+", stripped)
                or re.fullmatch(r"1[3-9]\d{9}", stripped)
                or re.fullmatch(r"(?:\+?86[-\s]?)?1[3-9]\d{9}", stripped)
            )

            if is_pure_noise or is_bare_contact:
                # 联系方式行：不丢弃，收集起来稍后插回文本
                contact_lines.append(stripped)
                continue

            if not stripped:
                continue

            # 判断这行是否包含课程碎片
            # 必须含课程相关后缀词，不能只是任意中文+·
            has_course_keyword = bool(re.search(
                r"(程序设计|系统|原理|分析|设计|算法|学|论|基础|技术|⼯程|密码|UNIX|LINUX|⽹络|视觉|智能|嵌⼊|汇编|软件)",
                stripped
            ))
            is_course_fragment = (
                ("·" in stripped and has_course_keyword)
                or (has_course_keyword and len(stripped) < 60)
            )

            if is_course_fragment:
                # 提取行内混入的邮箱/电话，保存到contact_lines
                email_m = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", stripped)
                phone_m = re.search(r"1[3-9]\d{9}", stripped)
                if email_m:
                    contact_lines.append("邮箱 " + email_m.group(0))
                if phone_m:
                    contact_lines.append("电话 " + phone_m.group(0))
                # 去掉行内混入的邮箱/电话片段
                clean = re.sub(r"(邮箱|电话|手机)\s*\S+@\S+", "", stripped)
                clean = re.sub(r"(邮箱|电话|手机)\s*1[3-9]\d{9}", "", clean)
                clean = re.sub(r"\S+@\S+\.\S+", "", clean)  # 去掉邮箱
                clean = re.sub(r"1[3-9]\d{9}", "", clean)  # 去掉手机号
                clean = clean.strip()
                if clean:
                    course_buffer.append(clean)
                continue
            else:
                # 非课程、非噪音行 → 课程段落结束
                in_courses = False
                if course_buffer:
                    merged = " ".join(course_buffer)
                    merged = re.sub(r"\s*·\s*$", "", merged)
                    merged = re.sub(r"·\s*·", "·", merged)
                    result_lines.append(merged)
                    course_buffer = []
                result_lines.append(line)
                continue

        result_lines.append(line)

    # 如果文件结尾还在课程段落中
    if course_buffer:
        merged = " ".join(course_buffer)
        merged = re.sub(r"\s*·\s*$", "", merged)
        merged = re.sub(r"·\s*·", "·", merged)
        result_lines.append(merged)

    # 把被课程段落截断的联系方式行插回文本顶部（姓名之后）
    if contact_lines:
        # 找到第一行非空行（通常是姓名），联系方式插在它后面
        insert_pos = 0
        for i, l in enumerate(result_lines):
            if l.strip():
                insert_pos = i + 1
                break
        for j, cl in enumerate(contact_lines):
            result_lines.insert(insert_pos + j, cl)

    return "\n".join(result_lines)


# ===== 姓名 =====

_NAME_BLACKLIST = {
    "个人简历", "求职简历", "基本信息", "个人信息", "联系方式", "求职意向",
    "教育背景", "教育经历", "工作经历", "项目经历", "项目经验", "技能特长",
    "自我评价", "个人简介", "简历", "求职信", "应聘", "实习生", "专业技能",
    "个人技能", "在校经历", "荣誉证书", "个人评价", "兴趣爱好",
}


def _ensure_name(llm_name: str, text: str) -> str:
    """校验LLM给的姓名，不对就正则提取"""
    name = (llm_name or "").strip()
    if _is_valid_name(name):
        return name
    return _regex_name(text)


def _is_valid_name(name: str) -> bool:
    """判断是否像合法中文姓名"""
    if not name or len(name) < 2 or len(name) > 4:
        return False
    if name in _NAME_BLACKLIST:
        return False
    if "@" in name or re.search(r"\d", name):
        return False
    # 必须是纯汉字（或·分隔的少数民族名）
    if not re.fullmatch(r"[\u4e00-\u9fff·]{2,6}", name):
        return False
    return True


def _regex_name(text: str) -> str:
    """多策略提取中文姓名"""
    # 策略1：明确标记 "姓名：XX" / "Name: XX"
    m = re.search(r"(?:姓\s*名|Name)[：:\s]+([^\s,，、/]{2,4})", text, re.IGNORECASE)
    if m and _is_valid_name(m.group(1)):
        return m.group(1)

    # 策略2：前5行找纯汉字短词
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:5]:
        # 去掉分隔符
        clean = re.sub(r"[|｜\-\s\t]", "", line)
        if _is_valid_name(clean):
            return clean
        # 行内可能有 "张三 | 13800001111" 这种
        parts = re.split(r"[|｜\s]{2,}", line)
        for part in parts:
            part = part.strip()
            if _is_valid_name(part):
                return part

    # 策略3：全文找"姓名"附近
    m = re.search(r"[\u4e00-\u9fff]{2,3}(?=\s*[|｜]\s*\d{11})", text)
    if m and _is_valid_name(m.group(0)):
        return m.group(0)

    return ""


def _regex_email(text: str) -> str:
    m = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return m.group(0) if m else ""


def _regex_phone(text: str) -> str:
    m = re.search(r"1[3-9]\d{9}", text)
    return m.group(0) if m else ""


# ===== 语言能力 =====

_LANG_NAMES = r"(?:英语|英文|English|法语|法文|French|日语|日文|Japanese|韩语|韩文|Korean|德语|German|西班牙语|Spanish|俄语|Russian|中文|普通话|Mandarin|粤语|Cantonese)"
_LANG_CERTS = r"(?:CET-?\d|TEM-?\d|IELTS|TOEFL|TOEIC|DELF\s*[A-C]\d|DALF\s*[A-C]\d|JLPT\s*N\d|J-TEST|HSK\s*\d|BULATS|BEC\s*(?:初级|中级|高级)?|~?[A-C]\d)"


def _regex_languages(text: str) -> list[str]:
    """从文本中提取语言能力信息"""
    results = []
    lines = text.split("\n")

    # 策略1：找"语言能力"段落，提取语言+证书
    in_lang_section = False
    current_lang = ""

    for line in lines:
        stripped = line.strip()

        # 进入语言能力段落
        if re.search(r"语[⾔言]能[⼒力]|Languages?", stripped) and len(stripped) < 15:
            in_lang_section = True
            continue

        # 离开段落
        if in_lang_section and re.match(
            r"^(?:教育|工作|实习|项目|技能|自我|荣誉|培训|开源|技术)", stripped
        ) and len(stripped) < 10:
            in_lang_section = False
            continue

        if not in_lang_section:
            # 策略2：全文扫描 "英语 CET-6" 这种同行格式
            m = re.search(
                rf"({_LANG_NAMES})\s*[:：]?\s*({_LANG_CERTS}[^\n,，]{{0,20}})",
                stripped, re.IGNORECASE
            )
            if m:
                entry = f"{m.group(1)} {m.group(2).strip()}"
                if entry not in results:
                    results.append(entry)
            continue

        # 在语言能力段落内
        # 检测语言名行（如 "英语"、"法语"）
        lang_m = re.match(rf"^({_LANG_NAMES})\s*(.*)", stripped, re.IGNORECASE)
        if lang_m:
            current_lang = lang_m.group(1)
            rest = lang_m.group(2).strip()
            # 同行有证书信息（排除纯日期范围如 "2018 至 2023"）
            if rest and re.search(r"[A-Z]", rest) and not re.fullmatch(r"20\d{2}\s*[至\-–—~]\s*20\d{2}", rest):
                entry = f"{current_lang} {rest}"
                if entry not in results:
                    results.append(entry)
                current_lang = ""
            continue

        # 检测证书行（如 "CET-6 / ~C1"、"DELF B1"）
        cert_m = re.search(rf"({_LANG_CERTS}(?:\s*/\s*~?[A-C]\d)?(?:（[^）]{{0,20}}）)?)", stripped, re.IGNORECASE)
        if cert_m and current_lang:
            entry = f"{current_lang} {cert_m.group(1).strip()}"
            if entry not in results:
                results.append(entry)
            current_lang = ""
        elif cert_m and not current_lang:
            # 没有明确语言名，尝试从上下文推断
            cert_text = cert_m.group(1).strip()
            if "CET" in cert_text or "TEM" in cert_text or "IELTS" in cert_text or "TOEFL" in cert_text:
                entry = f"英语 {cert_text}"
            elif "DELF" in cert_text or "DALF" in cert_text:
                entry = f"法语 {cert_text}"
            elif "JLPT" in cert_text:
                entry = f"日语 {cert_text}"
            else:
                entry = cert_text
            if entry not in results:
                results.append(entry)

    return results[:5]


# ===== 教育经历 =====

def _regex_education(text: str) -> list[dict]:
    """基于段落上下文提取教育经历"""
    results = []
    lines = text.split("\n")
    in_edu = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        # 进入教育段落
        if re.search(r"教育[背经][景历]", stripped) and len(stripped) < 10:
            in_edu = True
            continue
        # 离开教育段落（遇到下一个段落标题）
        if in_edu and re.search(r"(?:工作|实习|项目|技能|自我)[经特]", stripped) and len(stripped) < 10:
            in_edu = False
            continue

        if in_edu or not results:
            # 找大学名
            uni_match = re.search(r"([\u4e00-\u9fff]+(?:大学|学院))", stripped)
            if not uni_match:
                uni_match = re.search(r"([A-Za-z\s]+(?:University|College|Institute))", stripped)
            if uni_match:
                school = uni_match.group(1)
                # 同行提取其他信息
                degree_m = re.search(r"(博士|硕士|研究生|本科|学士|大专|专科)", stripped)
                major_m = re.search(r"(?:专业[：:\s]*)?([\u4e00-\u9fff]{2,10}(?:工程|学|技术|管理|设计|科学))", stripped)
                dates = re.findall(r"(20\d{2})[./\-年](\d{1,2})", stripped)
                start_date = f"{dates[0][0]}.{dates[0][1].zfill(2)}" if len(dates) >= 1 else ""
                end_date = f"{dates[1][0]}.{dates[1][1].zfill(2)}" if len(dates) >= 2 else ""

                # 如果同行没找到专业，看下一行
                major = ""
                if major_m:
                    major = major_m.group(1)
                elif i + 1 < len(lines):
                    next_l = lines[i + 1].strip()
                    major_m2 = re.search(r"([\u4e00-\u9fff]{2,10}(?:工程|学|技术|管理|设计|科学))", next_l)
                    if major_m2:
                        major = major_m2.group(1)

                results.append({
                    "school": school,
                    "major": major,
                    "degree": degree_m.group(1) if degree_m else "",
                    "start_date": start_date,
                    "end_date": end_date,
                })

    # 如果段落式没找到，全文扫描大学名
    if not results:
        for m in re.finditer(r"([\u4e00-\u9fff]+(?:大学|学院))", text):
            school = m.group(1)
            # 取大学名前后50字符作为上下文
            ctx = text[max(0, m.start() - 30):m.end() + 50]
            degree_m = re.search(r"(博士|硕士|研究生|本科|学士|大专|专科)", ctx)
            major_m = re.search(r"([\u4e00-\u9fff]{2,10}(?:工程|学|技术|管理|设计|科学))", ctx)
            dates = re.findall(r"(20\d{2})[./\-年](\d{1,2})", ctx)
            results.append({
                "school": school,
                "major": major_m.group(1) if major_m else "",
                "degree": degree_m.group(1) if degree_m else "",
                "start_date": f"{dates[0][0]}.{dates[0][1].zfill(2)}" if len(dates) >= 1 else "",
                "end_date": f"{dates[1][0]}.{dates[1][1].zfill(2)}" if len(dates) >= 2 else "",
            })
            if len(results) >= 3:
                break

    return results


# ===== 工作/实习经历 =====

_SECTION_HEADERS = re.compile(
    r"^(?:工作经[历验]|实习经[历验]|工作/实习|实践经[历验]|校园经[历验])$",
)


def _regex_experiences(text: str) -> list[dict]:
    """提取工作/实习经历"""
    results = []
    lines = text.split("\n")
    in_section = False
    current = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 进入经历段落
        if _SECTION_HEADERS.match(stripped):
            in_section = True
            continue
        # 离开（遇到下一个段落标题）
        if in_section and re.match(r"^(?:教育|项目|技能|自我|荣誉)", stripped) and len(stripped) < 10:
            in_section = False
            if current:
                results.append(current)
                current = None
            continue

        if not in_section:
            continue

        # 检测公司行（含公司关键词或含日期范围）
        has_date = re.search(r"20\d{2}[./\-]\d{1,2}", stripped)
        has_company = any(kw in stripped for kw in [
            "公司", "集团", "科技", "网络", "传媒", "有限", "工作室",
            "Co.", "Ltd", "Inc", "Corp",
        ])

        if has_company or (has_date and len(stripped) < 40):
            if current:
                results.append(current)
            # 提取日期
            duration = ""
            dur_m = re.search(
                r"(20\d{2}[./\-]\d{1,2}\s*[-–—~至]\s*(?:20\d{2}[./\-]\d{1,2}|至今|今))",
                stripped
            )
            if dur_m:
                duration = dur_m.group(1)
            company = stripped.replace(duration, "").strip(" |｜-—·")
            current = {"company": company[:30], "position": "", "duration": duration, "description": ""}
        elif current:
            # 第二行通常是职位
            if not current["position"] and len(stripped) < 25:
                current["position"] = stripped
            else:
                # 后续是描述
                if current["description"]:
                    current["description"] += "；" + stripped[:60]
                else:
                    current["description"] = stripped[:80]

    if current:
        results.append(current)
    return results[:5]


# ===== 项目经历 =====

def _regex_projects(text: str) -> list[dict]:
    """提取项目经历"""
    results = []
    lines = text.split("\n")
    in_section = False
    current = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if re.match(r"^项目[经]?[历验]?$", stripped):
            in_section = True
            continue
        if in_section and re.match(r"^(?:教育|工作|实习|技能|自我|荣誉)", stripped) and len(stripped) < 10:
            in_section = False
            if current:
                results.append(current)
                current = None
            continue

        if not in_section:
            continue

        # 技术栈行
        tech_m = re.search(r"(?:技术栈|Tech|使用技术|开发工具)[：:\s]*(.+)", stripped, re.IGNORECASE)
        if tech_m and current:
            current["tech_stack"] = tech_m.group(1).strip()[:80]
            continue

        # 角色行
        role_m = re.search(r"(?:角色|担任|职位)[：:\s]*(.+)", stripped)
        if role_m and current:
            current["role"] = role_m.group(1).strip()[:15]
            continue

        # 新项目标题（短行，不以标点开头）
        if len(stripped) < 30 and not stripped.startswith(("-", "·", "•", "–", "·")):
            if current:
                results.append(current)
            current = {"name": stripped, "role": "", "description": "", "tech_stack": ""}
        elif current:
            if not current["description"]:
                current["description"] = stripped[:100]

    if current:
        results.append(current)
    return results[:5]


# ===== 技能 =====

_KNOWN_TECHS = [
    "Python", "Java", "JavaScript", "TypeScript", "Go", "C++", "C#", "Rust", "C",
    "Vue", "Vue3", "React", "Angular", "FastAPI", "Flask", "Django", "Spring",
    "MySQL", "Redis", "MongoDB", "PostgreSQL", "SQL", "SQLite",
    "Docker", "Kubernetes", "K8s", "Linux", "Git", "GitHub",
    "HTML", "CSS", "Node.js", "Webpack", "Vite", "Tailwind",
    "PyTorch", "TensorFlow", "LangChain", "LangGraph", "Ollama", "ChromaDB", "Chroma",
    "Figma", "Axure", "Photoshop", "Excel", "PPT", "Word",
    "pandas", "numpy", "Spark", "Hive", "Tableau", "PowerBI",
    "微信小程序", "uni-app", "Flutter", "Android", "iOS",
    "Playwright", "Selenium", "YOLO", "RAG", "Agent", "Vibe Coding",
    "Server酱", "CAS", "Cookie", "API", "CLI",
    "嵌入式", "机器学习", "计算机视觉", "强化学习", "深度学习",
    "NLP", "CV", "LLM", "GPT", "Copilot", "Codex",
]


def _regex_skills(text: str) -> list[str]:
    """从文本中匹配已知技术词 + 技术栈段落提取"""
    found = []
    text_lower = text.lower()

    # 匹配已知技术词（短名用词边界避免误匹配）
    for tech in _KNOWN_TECHS:
        if len(tech) <= 2:
            # 短名（C, Go, CV等）需要词边界
            if re.search(r"(?<![a-zA-Z])" + re.escape(tech.lower()) + r"(?![a-zA-Z])", text_lower):
                found.append(tech)
        else:
            if tech.lower() in text_lower:
                found.append(tech)

    # 提取"技术栈"段落（后续几行都是技能词）
    lines = text.split("\n")
    in_tech_section = False
    tech_section_lines = 0
    for line in lines:
        stripped = line.strip()
        if re.match(r"^技术栈$", stripped) or re.match(r"^(?:技能|技术|Skills)[：:\s]*$", stripped, re.IGNORECASE):
            in_tech_section = True
            tech_section_lines = 0
            continue
        if in_tech_section:
            # 遇到下一个段落标题就停（允许标题带括号注释，如"开源项目（GitHub仓库）"）
            if re.match(r"^(?:教育|工作|实习|项目|自我|语言|培训|开源|基本|期望|荣誉)", stripped) and len(stripped) < 25:
                in_tech_section = False
                continue
            # 跳过"仓库:"行（GitHub仓库路径，不是技能）
            if re.match(r"^\s*仓库", stripped) or "/" in stripped and re.search(r"[A-Z][a-z]+/[A-Z]", stripped):
                continue
            if not stripped:
                continue
            tech_section_lines += 1
            if tech_section_lines > 6:
                in_tech_section = False
                continue
            # 先从行中提取已知多词技术（如 Vibe Coding）
            remaining = stripped
            for tech in _KNOWN_TECHS:
                if " " in tech and tech.lower() in remaining.lower():
                    if tech not in found:
                        found.append(tech)
                    remaining = re.sub(re.escape(tech), " ", remaining, flags=re.IGNORECASE)
            # 再按空格和分隔符切分剩余（技术栈行是空格分隔的）
            items = re.split(r"[,，、/|·\s]+", remaining)
            for item in items:
                item = item.strip()
                if item and len(item) <= 15 and item not in found:
                    found.append(item)

    # 扫描项目技术栈行（含 · 分隔的多个技术词）
    for line in lines:
        stripped = line.strip()
        # 跳过课程行（由 _courses_to_skills 单独处理）
        if re.search(r"核[心⼼]课程|课程", stripped):
            continue
        if "·" in stripped and stripped.count("·") >= 2:
            items = re.split(r"[·]+", stripped)
            tech_count = sum(1 for it in items if it.strip().lower() in [t.lower() for t in _KNOWN_TECHS])
            if tech_count >= 2:  # 至少2个已知技术词才算技术栈行
                for item in items:
                    item = item.strip()
                    if item and len(item) <= 15 and item not in found:
                        found.append(item)

    # 也尝试提取"技能：XXX, YYY"单行格式
    m = re.search(r"(?:技能|Skills)[：:\s]+(.+)", text, re.IGNORECASE)
    if m:
        items = re.split(r"[,，、/|·]+", m.group(1))
        for item in items:
            item = item.strip()
            if item and len(item) <= 15 and item not in found:
                found.append(item)

    return _filter_skills(found)


# 非技能黑名单模式
_SKILL_BLACKLIST_PATTERNS = [
    r"\d{4}",            # 含年份数字（2027届、2022.07等）
    r"^[‒\-–—~·•]+$",   # 纯分隔符
    r"大学|学院|University",  # 校名
    r"本科|硕士|博士|大专|学士",  # 学历
    r"中国|法国|美国|英国|德国",  # 国家名
    r"GPA|绩点",         # 成绩
    r"年龄|岁",          # 年龄
    r"邮箱|电话|手机",    # 联系方式标签
    r"至今|⾄今",        # 时间词
    r"课程|学[年分期]",   # 学术词
    r"提供|负责|覆盖|沟通|编制|协助|优化|推进|定位|安装|调试",  # 动词（句子碎片）
    r"团队|科研|国际|一线|⼀线",  # 描述词
    r"故障|排查|⼯单|响应|解决时间",  # 工作描述
    r"运营维护|技术支持|技术⽀持",  # 岗位描述
    r"项[⽬目]经[历验]|教育[背经][景历]|实习经[历验]|⾃我评价|语[⾔言]能[⼒力]|开源项[⽬目]|培训经[历验]",  # 段落标题
    r"仓库|GitHub|github",  # GitHub仓库相关
    r"^[\w-]+/[\w-]+",   # 路径格式（User/Repo）
    r"Watcher|Cloer|JNU",  # 仓库路径碎片
]

# 已知的中文技术词白名单（纯汉字但确实是技能）
_CHINESE_TECH_WHITELIST = {
    "嵌入式", "嵌⼊式", "机器学习", "计算机视觉", "强化学习", "深度学习",
    "自然语言处理", "数据分析", "数据挖掘", "微信小程序", "人工智能",
    "操作系统", "计算机网络", "数据库", "算法", "前端", "后端",
}

# 非技术的常见英文短词（不应作为技能）
_NON_TECH_ENGLISH = {
    "it", "service", "desk", "the", "and", "for", "with", "from",
    "this", "that", "have", "has", "was", "were", "are", "been",
    "can", "will", "would", "could", "should", "may", "might",
    "not", "but", "or", "if", "then", "than", "so", "as", "at",
    "in", "on", "to", "of", "by", "up", "out", "no", "yes",
    "all", "any", "each", "every", "both", "few", "more", "most",
    "other", "some", "such", "only", "own", "same", "too", "very",
    "just", "also", "now", "here", "there", "when", "where", "how",
    "what", "which", "who", "whom", "why", "because", "about",
    "into", "through", "during", "before", "after", "above", "below",
    "between", "under", "over", "again", "further", "once",
}


def _filter_skills(skills: list[str]) -> list[str]:
    """白名单式过滤：只保留看起来像技术/工具名的项"""
    known_lower = {t.lower() for t in _KNOWN_TECHS}
    result = []
    for s in skills:
        s = s.strip().strip("·•-–— ")
        if not s or len(s) < 1 or len(s) > 20:
            continue
        # 黑名单正则
        if any(re.search(p, s) for p in _SKILL_BLACKLIST_PATTERNS):
            continue
        # 已知技术词直接通过
        if s.lower() in known_lower:
            if s not in result:
                result.append(s)
            continue
        # 中文技术词白名单
        if s in _CHINESE_TECH_WHITELIST:
            if s not in result:
                result.append(s)
            continue
        # 纯汉字：只允许2-4字且在白名单中，否则拒绝
        if re.fullmatch(r"[\u4e00-\u9fff\u3400-\u4dbf]+", s):
            continue  # 不在白名单的纯汉字一律拒绝
        # 英文短词黑名单
        if s.lower() in _NON_TECH_ENGLISH:
            continue
        # 纯英文/数字/符号：必须看起来像技术词
        # 技术词特征：含大写字母、含数字、含特殊字符(+/#/.)、或<=10字符的英文
        if re.fullmatch(r"[a-zA-Z0-9+#./\- ]+", s):
            # 单个常见英文单词（非技术）拒绝
            if re.fullmatch(r"[a-zA-Z]+", s) and len(s) <= 6 and s.lower() in _NON_TECH_ENGLISH:
                continue
            # 看起来像技术词（含大写/数字/特殊字符，或短英文）
            if s not in result:
                result.append(s)
            continue
        # 其他情况（混合中英文等）：只保留短的
        if len(s) <= 8 and s not in result:
            result.append(s)
    return result


# ===== 相似技能合并 =====

# 技能合并组（有序列表）：同一组内的技能如果出现多个，合并为 "A/B/C" 形式
_SKILL_MERGE_GROUPS = [
    ["Vue", "Vue3", "Vue2"],
    ["SQL", "MySQL", "PostgreSQL", "SQLite"],
    ["LangChain", "LangGraph"],
    ["React", "React Native"],
    ["JavaScript", "TypeScript", "JS", "TS"],
    ["Kubernetes", "K8s"],
    ["Node.js", "Node"],
    ["PyTorch", "TensorFlow"],
    ["HTML", "CSS"],
    ["C", "C++", "C#"],
    ["机器学习", "深度学习"],
    ["NLP", "自然语言处理"],
    ["CV", "计算机视觉"],
    ["LLM", "GPT"],
    ["Chroma", "ChromaDB"],
    ["数据分析", "数据挖掘"],
    ["微信小程序", "uni-app"],
]


def _merge_similar_skills(skills: list[str]) -> list[str]:
    """将相似/同族技能合并为 A/B 形式，减少冗余标签"""
    # 建立 lowercase → 原始名 映射
    lower_map = {}
    for s in skills:
        lower_map[s.lower()] = s

    merged_set = set()  # 已被合并掉的技能（不再单独出现）
    result = []
    used_groups = set()

    for group in _SKILL_MERGE_GROUPS:
        # 找当前技能列表中属于该组的项
        matched = [s for s in skills if s in group or s.lower() in {g.lower() for g in group}]
        if len(matched) >= 2:
            # 多个同族技能 → 合并
            # 按组内顺序排列
            group_lower = [g.lower() for g in group]
            matched_sorted = sorted(matched, key=lambda x: group_lower.index(x.lower()) if x.lower() in group_lower else 99)
            merged_label = "/".join(matched_sorted)
            result.append(merged_label)
            for m in matched:
                merged_set.add(m)
            used_groups.add(id(group))
        elif len(matched) == 1:
            # 只有一个，不合并，保留原样（稍后统一添加）
            pass

    # 添加未被合并的技能
    for s in skills:
        if s not in merged_set and s not in result:
            result.append(s)

    return result


# ===== 课程→技能推断 =====

# 课程名 → 推断出的技能/知识标签（只保留JD中会出现的）
COURSE_SKILL_MAP = {
    "操作系统": ["操作系统", "Linux"],
    "操作系统原理": ["操作系统", "Linux"],
    "计算机网络": ["计算机网络", "TCP/IP", "HTTP"],
    "计算机⽹络": ["计算机网络", "TCP/IP", "HTTP"],
    "数据库系统原理": ["数据库", "SQL", "MySQL"],
    "数据库": ["数据库", "SQL", "MySQL"],
    "算法分析与设计": ["算法", "数据结构"],
    "算法": ["算法", "数据结构"],
    "数据结构": ["数据结构", "算法"],
    "机器学习": ["机器学习", "Python", "scikit-learn"],
    "强化学习": ["强化学习", "Python"],
    "强化学习与最优控制": ["强化学习", "Python"],
    "计算机视觉": ["计算机视觉", "OpenCV", "CNN"],
    "深度学习": ["深度学习", "PyTorch", "CNN"],
    "自然语言处理": ["NLP", "Transformer"],
    "Python程序设计": ["Python"],
    "Python": ["Python"],
    "C++程序设计": ["C++"],
    "高级语言程序设计": ["C"],
    "⾼级语⾔程序设计": ["C"],
    "汇编语言程序设计": ["汇编语言"],
    "汇编语⾔程序设计": ["汇编语言"],
    "软件工程": ["软件工程", "设计模式"],
    "软件⼯程": ["软件工程", "设计模式"],
    "嵌入式系统": ["嵌入式", "C", "RTOS"],
    "嵌⼊式系统": ["嵌入式", "C", "RTOS"],
    "UNIX/LINUX使用": ["Linux", "Shell"],
    "UNIX/LINUX使⽤": ["Linux", "Shell"],
    "计算机密码学": ["密码学", "信息安全"],
    "人工智能": ["人工智能", "机器学习"],
    "编译原理": ["编译原理"],
    "计算机组成原理": ["计算机组成"],
    "Web开发": ["HTML", "CSS", "JavaScript", "HTTP"],
    "移动应用开发": ["Android", "iOS"],
    "云计算": ["云计算", "分布式系统"],
    "大数据": ["大数据", "Spark", "Hadoop"],
    "信息安全": ["信息安全", "网络安全"],
    "网络安全": ["网络安全"],
}


def _courses_to_skills(text: str) -> list[str]:
    """从核心课程行提取课程名，推断对应技能"""
    # 找到核心课程行
    course_line = ""
    for line in text.split("\n"):
        if re.search(r"核[心⼼]课程", line):
            course_line = line
            break

    if not course_line:
        return []

    # 去掉"核心课程："前缀
    course_line = re.sub(r"^.*?核[心⼼]课程[：:\s]*", "", course_line)

    # 按 · 分隔提取课程名
    courses = re.split(r"[·•]+", course_line)
    courses = [c.strip() for c in courses if c.strip()]

    # 映射课程→技能
    inferred = []
    for course in courses:
        # 精确匹配
        if course in COURSE_SKILL_MAP:
            for skill in COURSE_SKILL_MAP[course]:
                if skill not in inferred:
                    inferred.append(skill)
        else:
            # 模糊匹配：课程名包含关键词
            for key, skills in COURSE_SKILL_MAP.items():
                if key in course or course in key:
                    for skill in skills:
                        if skill not in inferred:
                            inferred.append(skill)
                    break

    return inferred
